#!/usr/bin/python
# -*- coding: utf-8 -*-
#
# Copyright (C) 2016 - 2017, asterodeia <d_wang_890227@outlook.com>


import logging
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt

from utils import *


def load_data(book, sheet, config) -> pd.DataFrame:
    logging.info('reading data from excel file')

    no_header = config['no_header']
    header_row_count = None if no_header else 0
    try:
        dataframe = pd.read_excel(book, sheet, header=header_row_count, dtype=str)
    except FileNotFoundError:
        logging.error('data source file: {} not found, please check the file name'.format(book))
        logging.info('exit, please check and rerun this script')
        exit(-1)

    except Exception as e:
        logging.error('unexpected error: {}'.format(e))
        exit(-1)

    else:

        if no_header or not config['use_column_name']:
            column_count = dataframe.shape[1]
            column_name = [number_2_char(i) for i in range(column_count)]
            dataframe.columns = column_name

        return dataframe


def process_data(data, config):
    logging.info('starting to process data')

    if 'choice_col' in config:
        needed_data_column = config['question_col'] + config['answer_col'] + config['choice_col']
    else:
        needed_data_column = config['question_col'] + config['answer_col']

    needed_data = data[needed_data_column]

    if config['trim']:
        logging.info('\ttrimming data')
        pattern = re.compile(r'\n(?=\S)')
        needed_data = needed_data.applymap(lambda s: re.sub(pattern, '', s))

    if config.get('add_conventional_letter'):

        logging.info('\tadding conventional letter to choices')
        for i in range(len(config['choice_col'])):
            conventional_letter = number_2_char(i)

            def add_conventional_letter(s):
                return s if s == 'nan' else '{}. {}'.format(conventional_letter, s)

            # don't know if there is a better way to update just one column of dataframe
            ser = needed_data[config['choice_col'][i]]
            ser = ser.apply(add_conventional_letter)

            needed_data[config['choice_col'][i]] = ser

    if config.get('change_answer_letter'):
        # this is for objective questions, so only one column of answer exist
        logging.info('\tchanging answer number to letters')
        answer_col = needed_data[config['answer_col'][0]]

        ser = answer_col.apply(lambda s: ''.join([number_2_char(int(ch), 1) for ch in s]))
        needed_data[config['answer_col'][0]] = ser

    logging.info('\treformatting data')
    ret = []
    for index, row in needed_data.iterrows():
        entry = dict(question=[], choices=[], answer=[])

        # question
        for q in row[config['question_col']]:
            entry['question'].append(q)

        # choices
        if 'choice_col' in config:
            for c in row[config['choice_col']]:
                # TODO: 'nan' values should be handle more carefully
                if c != 'nan':
                    entry['choices'].append(c)

        # answer
        for a in row[config['answer_col']]:
            entry['answer'].append(a)

        ret.append(entry)

    logging.info('done processing data')
    # return formatted question data
    return ret


# documnet builder
class DocumentWriter(object):
    def __init__(self, template):
        self._document = None
        self._template = template

    def create_document(self):

        if self._template and not check_file_exist(self._template):
            logging.warning('specified document template does not exist, use default template instead')
            self._template = None

        try:
            self._document = Document(self._template)

        # TODO: NOT CORRECT RIGHT NOW, CHANGE LATER
        except ValueError as e:
            logging.error(e)
        else:
            if not self._template:
                self.document.styles.add_style('答案', WD_STYLE_TYPE.PARAGRAPH)
                font = self.document.styles['Normal'].font
                font.size = Pt(12)

    @property
    def document(self):
        return self._document

    def get_document_params(self, unit='twips') -> dict:
        """ return a dict contains object of length of the last section parameters """

        assert (unit in ('cm', 'mm', 'twips', 'emu', 'inches', 'pt'))

        # here we use a small trick to avoid defining functions like:
        # if unit == 'cm':
        #   def func(x):
        #     return x.cm
        # elif: unit == 'mm':
        # ...
        func = eval('lambda x: x.{}'.format(unit))

        section = self.document.sections[-1]
        left_margin = func(section.left_margin)
        right_margin = func(section.right_margin)
        top_margin = func(section.top_margin)
        bottom_margin = func(section.bottom_margin)
        page_height = func(section.page_height)
        page_width = func(section.page_width)
        # 装订线
        gutter = section.gutter

        ret = {'margin': dict(top=top_margin, bottom=bottom_margin, left=left_margin, right=right_margin),
               'page_size': dict(height=page_height, width=page_width), 'gutter': gutter}

        return ret

    def write_title(self, title, level):
        assert 0 <= level <= 9
        self.document.add_heading(title, level)

    def write_question(self, question, style):
        for item in question:
            self.document.add_paragraph(item, style=style)

    def write_choices(self, choices, style):
        # this method is not interested by every derived class,
        # so doesn't need a default implementation
        pass

    def write_answer(self, answer, style):
        for item in answer:
            self.document.add_paragraph(item, style=style)


class ObjectiveQuestionWriter(DocumentWriter):
    def __init__(self, template=None):
        super().__init__(template)

    def _get_choice_tabstops_pos(self, choices: list, style):

        page_params = self.get_document_params('cm')
        style = self.document.styles[style]
        font_size = style.font.size.cm
        start_pos = 2 * font_size

        page_width = page_params['page_size']['width']
        available_width = page_width - page_params['margin']['left'] - page_params['margin']['right']
        single_choice_width = (available_width - start_pos - 1) / 4
        single_choice_length_4 = single_choice_width // font_size
        single_choice_length_2 = single_choice_width // font_size * 2

        max_ans_len = max(map(len, choices))

        if max_ans_len <= single_choice_length_4:
            tab_pos = [start_pos + x * single_choice_width for x in range(4)]
        elif max_ans_len <= single_choice_length_2:
            tab_pos = [start_pos + x * single_choice_width * 2 for x in range(2)]
        else:
            tab_pos = [start_pos]

        return tab_pos

    @staticmethod
    def _get_formatted_choice(choices: list, choice_count_per_line: int):
        formatted_choices = ['\t']
        current_item_count = choice_count_per_line

        for idx, choice in enumerate(choices, 1):
            formatted_choices.append(choice)
            current_item_count -= 1

            # decide string separator
            if idx != len(choices):
                if current_item_count != 0:
                    formatted_choices.append('\t')
                else:
                    formatted_choices.append('\n\t')
                    current_item_count = choice_count_per_line
            else:
                pass
                # don't need anything

        return ''.join(formatted_choices)

    def write_choices(self, choices: list, style: str):

        tab_pos = self._get_choice_tabstops_pos(choices, style)
        formatted_choice = self._get_formatted_choice(choices, len(tab_pos))

        p = self.document.add_paragraph(''.join(formatted_choice))
        pf = p.paragraph_format
        for pos in tab_pos:
            pf.tab_stops.add_tab_stop(Cm(pos))

    def write_answer(self, answer, style):

        # objective questions' answers are like 'A' or 'BCD',
        # so just covert the list of answer to a string by concatenate them
        if isinstance(answer, list):
            formatted_answer = concatenate(answer)
        else:
            formatted_answer = answer

        p = self.document.add_paragraph("\t答案：{}".format(formatted_answer), style=style)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


class TrueFalseWriter(DocumentWriter):
    def __init__(self, template=None):
        super().__init__(template)

    def write_answer(self, answer, style):
        # this type of questions should have only one answer
        assert len(answer) == 1
        p = self.document.add_paragraph("\t答案：" + answer[0], style='答案')
        p.paragraph_format.alignment = WD_TAB_ALIGNMENT.RIGHT


class SubjectiveQuestionWriter(DocumentWriter):
    def __init__(self, template=None):
        super().__init__(template)

    def write_answer(self, answer, style):
        self.document.add_paragraph("\t答案：")

        for item in answer:
            self.document.add_paragraph(item, style=style)


class QuestionRecompositor(object):
    # writer: for later use, now is None
    def __init__(self, config, writer):
        self._config = config
        self._writer = writer
        # this may change later
        self.data_processor = process_data

    def set_config(self, config):
        self._config = config

    def _load_data(self, file_name, sheet_name):
        data_frame = load_data(file_name, sheet_name, self._config.read)
        return data_frame

    # director
    def recompose(self, file_name, sheet_name, builder: DocumentWriter):
        data_frame = self._load_data(file_name, sheet_name)
        data_frame = self.data_processor(data_frame, self._config.process)

        builder.create_document()
        builder.write_title(self._config.write['title'], level=0)

        logging.info('writing questions to document, {} in total'.format(len(data_frame)))

        for entry in data_frame:
            builder.write_question(entry['question'], 'Heading 2')

            if entry['choices']:
                # this is an objective question
                builder.write_choices(entry['choices'], 'Normal')

            if not self._config.write['hide_answer']:
                builder.write_answer(entry['answer'], '答案')
