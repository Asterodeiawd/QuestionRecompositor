#!/usr/bin/python
# -*- coding: utf-8 -*-
#
# Copyright (C) 2016 - 2017, asterodeia <noachic.wang@gmail.com>

"""
usage:
./QuestionRecompositor.py -i in.xls(x) -o output(.docx) --no_title --conventional_letter
--change_answer_letter --separate_choices --separator '\s'
"""

import json
import logging
import re
from argparse import ArgumentParser
from os import getcwd
from os import path

import pandas as pd
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s %(levelname)s: %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')


class QuestionRecompositor(object):
    def __init__(self, config, writer):
        self.config = config
        self.writer = writer
        # this may change later
        self.data_processor = process_data

    def set_config(self, config):
        self.config = config

    def set_writer(self, writer):
        self.writer = writer

    def _load_data(self, file_name, sheet_name):
        data_frame = load_data(file_name, sheet_name, self.config.read)
        return data_frame

    def recompose(self, file_name, sheet_name, doc):
        data_frame = self._load_data(file_name, sheet_name)
        data_frame = self.data_processor(data_frame, self.config.process)
        self.writer(data_frame, doc, self.config.write)


def load_data(book, sheet, config) -> pd.DataFrame:
    logging.info('reading data from excel file')

    no_header = config['no_header']
    header_row_count = None if no_header else 0
    try:
        dataframe = pd.read_excel(book, sheet, header=header_row_count, dtype=str)
    except FileNotFoundError as e:
        logging.error('data source file: {} not found, please check the file name'.format(book))
        logging.info('exit, please check and rerun this script')
        exit(-1)

    except Exception as e:
        logging.error('unexpected error: {}'.format(e))
        exit(-1)

    else:

        if no_header:
            column_count = dataframe.shape[1]
            column_name = [number_2_char(i) for i in range(column_count)]
            dataframe.columns = column_name

        return dataframe


def shift_character(ch, i):
    return chr(ord(ch) + i)


def number_2_char(n, base=0):
    assert n >= base
    return chr(ord('A') + n - base)


def process_data(data, config):
    logging.info('starting to process data')

    if 'choice_col' in config:
        needed_data_column = config['question_col'] + config['answer_col'] + config['choice_col']
    else:
        needed_data_column = config['question_col'] + config['answer_col']

    needed_data = data[needed_data_column]

    if config['trim']:
        logging.info('\ttrimming data')
        pattern = re.compile(r'[\r\n]')
        needed_data = needed_data.applymap(lambda s: re.sub(pattern, '', s))

    # TODO: if not choice questions, skip this
    if 'add_conventional_letter' in config and config['add_conventional_letter']:

        logging.info('\tadding conventional letter to choices')
        for i in range(len(config['choice_col'])):
            conventional_letter = number_2_char(i)

            def add_conventional_letter(s):
                return s if s == 'nan' else '{}. {}'.format(conventional_letter, s)

            # don't know if there is a better way to update just one column of dataframe
            ser = needed_data[config['choice_col'][i]]
            ser = ser.apply(add_conventional_letter)

            needed_data[config['choice_col'][i]] = ser

    if 'change_answer_letter' in config and config['change_answer_letter']:
        # TODO: first column of answer, change this?
        logging.info('\tchanging answer number to letters')
        answer_col = needed_data[config['answer_col'][0]]

        # def change_answer_letter(s):
        # answer = [number_2_char(int(ch), 1) for ch in s]
        # return ''.join(answer)

        ser = answer_col.apply(lambda s: ''.join([number_2_char(int(ch), 1) for ch in s]))
        needed_data[config['answer_col'][0]] = ser

    logging.info('\treformatting data')
    ret = []
    for index, row in needed_data.iterrows():
        entry = dict(question=[], choices=[], answer=[])

        # question
        for q in row[config['question_col']]:
            entry['question'].append(q)

        # choices
        if 'choice_col' in config:
            for c in row[config['choice_col']]:
                # TODO: 'nan' values should be handle more carefully
                if c != 'nan':
                    entry['choices'].append(c)

        # answer
        for a in row[config['answer_col']]:
            entry['answer'].append(a)

        ret.append(entry)

    logging.info('done processing data')
    # return formatted question data
    return ret


def get_document_params(doc):
    section = doc.sections[-1]
    left_margin = section.left_margin.cm
    right_margin = section.right_margin.cm
    top_margin = section.top_margin.cm
    bottom_margin = section.bottom_margin.cm
    height = section.page_height.cm
    width = section.page_width.cm
    print(height, width, left_margin, right_margin, top_margin, bottom_margin)


def write_doc_title(doc: Document, config):
    doc.add_heading(config['title'], 0)
    # doc.paragraphs[0].delete()


def write_doc_sc(data, doc: Document, config):
    # add document title
    write_doc_title(doc, config)

    logging.info('writing questions to document, {} in total'.format(len(data)))

    for entry in data:

        # question
        for item in entry['question']:
            doc.add_paragraph(item, style='Heading 2')

        # choices:
        max_ans_len = max(map(len, entry['choices']))
        if max_ans_len <= 10:
            separator = ['\t', '\t']
            tab_pos = [1 + x * 4.5 for x in range(4)]

        elif max_ans_len <= 22:
            separator = ['\t', '\n\t']
            tab_pos = [1, 10]

        else:
            separator = ['\n\t', '\n\t']
            tab_pos = [1]

        choice_count = len(entry['choices'])
        choice_string = ['\t']
        for cnt, item in enumerate(entry['choices']):
            if cnt != choice_count - 1:
                choice_string += item + separator[cnt % 2]
            else:
                choice_string += item

        p = doc.add_paragraph(''.join(choice_string))
        pf = p.paragraph_format
        for pos in tab_pos:
            pf.tab_stops.add_tab_stop(Cm(pos))

        # answer
        if not config['hide_answer']:
            answer_string = ''.join(entry['answer'])
            p = doc.add_paragraph("\t答案：" + answer_string, style='答案')
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Cm(18), alignment=WD_TAB_ALIGNMENT.RIGHT)

    return doc


# multiple choice
def write_doc_mc(data, doc: Document, config):
    # add document title
    write_doc_title(doc, config)

    logging.info('writing questions to document, {} in total'.format(len(data)))

    for entry in data:

        # question
        for item in entry['question']:
            doc.add_paragraph(item, style='Heading 2')

        # choices:
        max_ans_len = max(map(len, entry['choices']))
        if max_ans_len <= 10:
            separater = ['\t', '\t']
            tab_pos = [1 + x * 4.5 for x in range(4)]

        elif max_ans_len <= 22:
            separater = ['\t', '\n\t']
            tab_pos = [1, 10]

        else:
            separater = ['\n\t', '\n\t']
            tab_pos = [1]

        choice_count = len(entry['choices'])
        choice_string = ['\t']
        for cnt, item in enumerate(entry['choices']):
            if cnt != choice_count - 1:
                choice_string += item + separater[cnt % 2]
            else:
                choice_string += item

        p = doc.add_paragraph(''.join(choice_string))
        pf = p.paragraph_format
        for pos in tab_pos:
            pf.tab_stops.add_tab_stop(Cm(pos))

        # answer
        if not config['hide_answer']:
            answer_string = ''.join(entry['answer'])
            p = doc.add_paragraph("\t答案：" + answer_string, style='答案')
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Cm(18), alignment=WD_TAB_ALIGNMENT.RIGHT)

    return doc


# true false questions
def write_doc_tf(data, doc: Document, config):
    # add document title
    write_doc_title(doc, config)

    logging.info('writing questions to document, {} in total'.format(len(data)))

    for entry in data:

        # question
        for item in entry['question']:
            doc.add_paragraph(item, style='Heading 2')

        # answer
        if not config['hide_answer']:
            answer_string = ''.join(entry['answer'])
            p = doc.add_paragraph("\t答案：" + answer_string, style='答案')
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Cm(18), alignment=WD_TAB_ALIGNMENT.RIGHT)

    return doc


def main():
    cwd = getcwd()
    source_path = path.join(cwd, 'data', r'线路架空线专业单选.xls')
    # source_path = r"C:\Users\asterodeia\Desktop\变压器新增题库汇总.xlsx"
    config_path = path.join(cwd, 'config.json')
    document_template_path = path.join(cwd, 'data', 'template.docx')

    configs = load_configs_from_file(config_path)
    c = configs[0]
    print(c.name)

    name = 'Sheet1'

    doc = Document(document_template_path)
    qr = QuestionRecompositor(c, write_doc_sc)
    qr.recompose(source_path, name, doc)
    doc.save(path.join(cwd, name))


def main2():
    parser = ArgumentParser(description='a python script to convert structured excel to word')
    parser.add_argument('-s', '--single-choice', metavar='filename', action='store',
                        help='source structured excel file')
    parser.add_argument('-m', '--multi-choice', metavar='filename', action='store',
                        help='target word file name (default: output.docx)')
    parser.add_argument('-t', '--true-false', action='store_true',
                        help="to specify the excel file don't have a title row")
    args = parser.parse_args()
    pass


class Config(object):
    def __init__(self, data):
        # configs should include 4 parts: name, read method, process method and write method
        # so here don't have to check before assign
        self.name = None
        self.read = None
        self.process = None
        self.write = None

        self.load(data)

    def load(self, data):
        self.name = data['name']
        self.read = data['read']
        self.process = data['process']
        self.write = data['write']

    def unify_process(self):
        if not self.process['use_column_name']:
            for col in ['question_col', 'answer_col', 'choice_col']:
                if isinstance(self.process[col], str):
                    self.process[col] = list(self.process[col])


def load_configs_from_file(filename: str) -> list:
    """
     helper function to load configs from file
    :param filename: file contains the config(s)
    :return: list of Config
    """
    configs = []

    with open(filename, mode='r', encoding='utf-8') as f:
        raw_configs = json.loads(f.read())

        # configs should be in one list, single config can omit the outter list
        if isinstance(raw_configs, list):
            for raw_config in raw_configs:
                configs.append(Config(raw_config))
        else:
            configs.append(Config(raw_configs))

    logging.info('unifying process part in config')
    for config in configs:
        config.unify_process()

    return configs


if __name__ == '__main__':
    # cwd = getcwd()
    # document_template_path = path.join(cwd, 'data', 'template.docx')
    # doc = Document(document_template_path)
    # get_document_params(doc)
    main()
